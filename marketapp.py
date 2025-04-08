from stock_utility_handler import StockAPI, StockAnalyzer
from ai_insights_handler import AIInsights
from fundamental_handler import get_fundamental_data

import streamlit as st
import os
import tempfile
from docx import Document
from docx.shared import Inches

# Initialize session state variables
if 'page' not in st.session_state:
    st.session_state.page = "page1"
    st.session_state.ticker = "RELIANCE"
    st.session_state.market = "BSE"
    st.session_state.analysis_type = "Technical"
    st.session_state.image_path = ""
    st.session_state.ai_insights = ""
    st.session_state.fundamentals = ""
    st.session_state.internal_results_available = False

# Page 1: Input Page
def page1():
    st.title('📈 Stock AI Agent')

    st.sidebar.header("ℹ️ About This AI-Powered Stock Analysis Bot")
    st.sidebar.write("""
    ### 📌 Overview  
    This AI-powered bot helps **investors, traders**, and **analysts** with insights from both technical charts and fundamental data.

    ### 🔍 What You Get  
    - ✅ Technical Charts + AI Insights  
    - ✅ Company Financials (Fundamentals)  
    - ✅ Market Support: BSE & NASDAQ  
    - ✅ Downloadable Reports (DOCX)

    🚀 Built with AI & Machine Learning
    """)

    col1, col2 = st.columns(2)
    with col1:
        st.session_state.ticker = st.text_input("🏷️ Enter Stock Ticker", value=st.session_state.ticker, key="ticker_input")
    with col2:
        st.session_state.market = st.selectbox("🌍 Select Market", ["BSE", "NASDAQ"], index=["BSE", "NASDAQ"].index(st.session_state.market), key="market_input")

    st.session_state.analysis_type = st.selectbox("📊 Select Analysis Type", ["Technical", "Fundamental", "Both"])

    st.markdown("---")

    if st.button('🚀 Submit'):
        st.session_state.page = "page2"
        st.session_state.internal_results_available = False
        st.rerun()

# Page 2: Analysis Page
def page2():
    st.title(f"📊 Analysis for {st.session_state.ticker} ({st.session_state.market})")

    stock = st.session_state.ticker
    market = st.session_state.market
    analysis_type = st.session_state.analysis_type

    if not st.session_state.internal_results_available:
        with st.spinner('🔍 Analyzing... Please wait...'):
            temp_dir = tempfile.gettempdir()
            image_path = os.path.join(temp_dir, f"{market}_{stock}.png")
            st.session_state.image_path = image_path

            try:
                stock_api_obj = StockAPI("1UJ6ACYM0P4MHORZ")
                stock_analyzer_obj = StockAnalyzer()
                ai_insights_obj = AIInsights("AIzaSyAVi1v80vt41mTjZED6BaMs5-74HKFkSk0")

                doc = Document()
                doc.add_heading(f"Stock Analysis for {stock} ({market})", level=1)

                if analysis_type in ["Technical", "Both"]:
                    market_data = stock_api_obj.get_stock_info(stock, market)
                    df = stock_analyzer_obj.json_to_dataframe(market_data, stock, market)
                    fib_levels = stock_analyzer_obj.calculate_fibonacci_levels(df)
                    stock_analyzer_obj.plot_stock_data(df, stock, market, image_path, fib_levels)

                    ai_response = ai_insights_obj.get_ai_insights(image_path, stock, market)
                    ai_text = "\n".join([part.text for candidate in ai_response.candidates for part in candidate.content.parts])
                    st.session_state.ai_insights = ai_text

                    doc.add_heading("💡 Technical Insights", level=2)
                    doc.add_paragraph(ai_text)
                    doc.add_picture(image_path, width=Inches(5))

                if analysis_type in ["Fundamental", "Both"]:
                    fundamentals = get_fundamental_data(stock, market)
                    st.session_state.fundamentals = fundamentals
                    doc.add_heading("📚 Fundamental Analysis", level=2)
                    for key, value in fundamentals.items():
                        doc.add_paragraph(f"{key}: {value}")

                doc_path = os.path.join(tempfile.gettempdir(), f"{stock}_{market}_{analysis_type}_analysis.docx")
                doc.save(doc_path)
                st.session_state.doc_path = doc_path
                st.session_state.internal_results_available = True

            except Exception as e:
                st.error(f"❌ An error occurred: {e}")
                return

    # Display Results
    if st.session_state.internal_results_available:
        if analysis_type in ["Technical", "Both"]:
            st.subheader("📈 Stock Chart")
            st.image(st.session_state.image_path, caption=f"{stock} Chart", use_container_width=True)
            st.subheader("💡 AI Insights")
            st.write(st.session_state.ai_insights)

        if analysis_type in ["Fundamental", "Both"]:
            st.subheader("📚 Fundamentals")
            for k, v in st.session_state.fundamentals.items():
                st.write(f"**{k}:** {v}")

        with open(st.session_state.doc_path, "rb") as file:
            st.download_button(
                label=f"📥 Download {analysis_type} Report (Docx)",
                data=file,
                file_name=os.path.basename(st.session_state.doc_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.markdown("---")
        if st.button("🔙 Back to Home"):
            st.session_state.page = "page1"
            st.session_state.internal_results_available = False
            st.rerun()

# Page Routing
if st.session_state.page == "page1":
    page1()
elif st.session_state.page == "page2":
    page2()
